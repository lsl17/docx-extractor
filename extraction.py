from docx import Document
from typing import Any, List, Dict, Union, Optional
import json
from xml.etree import ElementTree as ET


def remove_consecutive_duplicates(lst: List[Any]) -> List[Any]:
    """Remove consecutive duplicates in a list."""
    if not lst:
        return []

    new_lst = [lst[0]]
    for i in range(1, len(lst)):
        if lst[i] != lst[i-1]:
            new_lst.append(lst[i])
    return new_lst


def extract_text_from_headers_and_footers(document: Document) -> List[Dict[str, Any]]:
    """
    Extract all text from the headers and footers of the document, including tables.

    :param document: A loaded DOCX document.
    :return: List of extracted content items.
    """
    header_contents = []
    footer_contents = []

    def process_container(container, location: str) -> List[Dict[str, Any]]:
        """Process content within a header or footer."""
        local_texts = []
        local_tables = []

        for element in container._element:
            if element.tag.endswith('tbl'):
                for table in container.tables:
                    if table._element == element:
                        table_data = extract_table(table)
                        local_tables.append(
                            {"table": table_data, "location": location})
                        break
            elif element.tag.endswith('p'):
                paragraph = [
                    p for p in container.paragraphs if p._element == element][0]
                paragraph_text = ''.join(run.text for run in paragraph.runs)
                if paragraph_text:
                    paragraph_text = paragraph_text.strip()
                    if paragraph_text:
                        local_texts.append(
                            {"text": paragraph_text, "location": location})

        return local_texts + local_tables

    for section in document.sections:
        header = section.header
        header_contents.extend(process_container(header, "header"))
        footer = section.footer
        footer_contents.extend(process_container(footer, "footer"))

    # Ensuring headers come before footers in the contents list
    return header_contents + footer_contents


def extract_content_from_document(document: Document) -> List[Dict[str, Any]]:
    """
    Extract content from the main body, headers, footers, footnotes, and endnotes of a DOCX document.

    :param document: A loaded DOCX document.
    :return: List of extracted content items.
    """
    content_list = []

    for element in document.element.body:
        if element.tag.endswith('tbl'):
            for table in document.tables:
                if table._element == element:
                    table_data = extract_table(table)
                    content_list.append({"table": table_data})
                    break
        elif element.tag.endswith('p'):
            paragraph = [
                p for p in document.paragraphs if p._element == element][0]
            paragraph_text = ''.join(run.text for run in paragraph.runs)
            if paragraph_text:
                paragraph_text = paragraph_text.strip()
                if paragraph_text:
                    content_list.append(
                        {"text": paragraph_text, "location": "body"})

    header_footer_texts = extract_text_from_headers_and_footers(document)
    content_list.extend(header_footer_texts)

    notes_texts = extract_footnotes_from_xml(document)
    content_list.extend(notes_texts)

    return content_list


def extract_footnotes_from_xml(document: Document) -> List[Dict[str, Any]]:
    """Extract footnotes from the document's XML content."""
    footnotes = []
    footnotes_part = None
    for rel in document.part.rels.values():
        if "footnotes" in rel.reltype:
            footnotes_part = rel.target_part
            break

    if not footnotes_part:
        return footnotes

    root = ET.fromstring(footnotes_part.blob)
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }

    for footnote in root.findall('.//w:footnote', namespaces=namespaces):
        text = []
        for child in footnote.findall('.//w:t', namespaces=namespaces):
            if child.text:
                text.append(child.text)
        footnotes.append({'text': ' '.join(text).strip(), 'type': 'footnote'})

    return footnotes


def extract_content_from_cell(cell) -> List[Dict[str, Any]]:
    """Extract content from a cell, preserving the order of text and nested tables."""
    content_list = []

    for element in cell._element:
        if element.tag.endswith('tbl'):
            for table in cell.tables:
                if table._element == element:
                    table_data = extract_table(table)
                    content_list.append({"table": table_data})
                    break
        elif element.tag.endswith('p'):
            paragraph = [
                p for p in cell.paragraphs if p._element == element][0]
            paragraph_text = ''.join(run.text for run in paragraph.runs)
            if paragraph_text:
                paragraph_text = paragraph_text.strip()
                if paragraph_text:
                    content_list.append({"text": paragraph_text})

    return content_list


def extract_table(table) -> List[List[Dict[str, Any]]]:
    """Extract data from a table with consecutive duplicates removal."""
    table_data = []

    for row in table.rows:
        row_data = []

        for cell in row.cells:
            cell_content_list = extract_content_from_cell(cell)
            row_data.append(cell_content_list)

        row_data = remove_consecutive_duplicates(row_data)
        table_data.append(row_data)

    return table_data


def extract_content_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Extract content from a DOCX file.

    :param file_path: Path to the DOCX file.
    :return: List of extracted content items.
    """
    try:
        document = Document(file_path)
        return extract_content_from_document(document)
    except Exception as e:
        print(f"Error processing {file_path}. Reason: {str(e)}")
        return []


def save_to_json(contents: List[Dict[str, Any]], output_path: str) -> None:
    """Save extracted content to a JSON file."""
    try:
        with open(output_path, 'w', encoding='utf-8') as json_file:
            json.dump(contents, json_file, indent=4, ensure_ascii=False)
        print(f"Data saved to {output_path}")
    except Exception as e:
        print(f"Error saving to {output_path}. Reason: {str(e)}")


def process_docx_and_save_to_json(input_file_path: str, output_file_path: str) -> None:
    """
    Extract content from a DOCX file and save the extracted content to a JSON file.

    :param input_file_path: Path to the DOCX file.
    :param output_file_path: Path where the extracted content will be saved as JSON.
    """
    contents = extract_content_from_docx(input_file_path)
    save_to_json(contents, output_file_path)


# Usage example:
if __name__ == '__main__':
    file_path = "sources/KFS Paired document/KFS - CT European Fund_E.docx"
    output_path = "KFS - CT European Fund_E.json"
    process_docx_and_save_to_json(file_path, output_path)
